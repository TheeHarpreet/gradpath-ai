"""Safe, in-memory document parsing and deterministic CV export."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Flowable,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
ALLOWED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


class DocumentProcessingError(ValueError):
    """An uploaded document cannot be safely converted to useful text."""


def extract_document_text(filename: str, content: bytes) -> str:
    """Extract text in memory from a small allow-listed document."""

    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentProcessingError(
            "Use a TXT, Markdown, DOCX, or text-based PDF file."
        )
    if not content:
        raise DocumentProcessingError("The uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise DocumentProcessingError("The uploaded file is larger than 5 MB.")

    try:
        if extension in {".txt", ".md"}:
            text = content.decode("utf-8-sig")
        elif extension == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        else:
            reader = PdfReader(BytesIO(content))
            if reader.is_encrypted:
                raise DocumentProcessingError(
                    "Password-protected PDFs are not supported."
                )
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError(
            "The file could not be read. Check that it is not damaged."
        ) from exc

    normalized = "\n".join(
        line.rstrip() for line in text.replace("\x00", "").splitlines()
    ).strip()
    if len(normalized) < 20:
        raise DocumentProcessingError(
            "Not enough selectable text was found. Scanned images need OCR "
            "before upload."
        )
    return normalized


def _markdown_blocks(markdown: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            blocks.append(("space", ""))
        elif line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("### "):
            blocks.append(("subheading", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("heading", line[3:].strip()))
        elif re.match(r"^[-*]\s+", line):
            blocks.append(("bullet", re.sub(r"^[-*]\s+", "", line)))
        else:
            blocks.append(("body", re.sub(r"\*\*(.*?)\*\*", r"\1", line)))
    return blocks


def build_cv_docx(markdown: str) -> bytes:
    """Build an editable, single-column A4 CV suitable for ATS parsing."""

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(31, 42, 38)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for kind, text in _markdown_blocks(markdown):
        if kind == "space":
            continue
        paragraph = document.add_paragraph()
        if kind == "title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(19)
            run.font.color.rgb = RGBColor(20, 73, 62)
        elif kind == "heading":
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text.upper())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(20, 73, 62)
        elif kind == "subheading":
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 42, 38)
        elif kind == "bullet":
            paragraph.style = document.styles["List Bullet"]
            paragraph.paragraph_format.left_indent = Mm(5)
            paragraph.paragraph_format.first_line_indent = Mm(-3)
            paragraph.add_run(text)
        else:
            paragraph.add_run(text)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_cv_pdf(markdown: str) -> bytes:
    """Build a fixed-layout PDF matching the clean DOCX hierarchy."""

    output = BytesIO()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CvBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        textColor=HexColor("#1f2a26"),
        spaceAfter=4,
    )
    title = ParagraphStyle(
        "CvTitle",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=19,
        leading=23,
        alignment=TA_CENTER,
        textColor=HexColor("#14493e"),
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "CvHeading",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=HexColor("#14493e"),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True,
    )
    subheading = ParagraphStyle(
        "CvSubheading",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        textColor=HexColor("#1f2a26"),
        spaceBefore=5,
        spaceAfter=2,
        keepWithNext=True,
    )
    bullet = ParagraphStyle(
        "CvBullet",
        parent=body,
        leftIndent=5 * mm,
        firstLineIndent=-3 * mm,
        bulletIndent=2 * mm,
    )
    story: list[Flowable] = []
    blocks = _markdown_blocks(markdown)
    index = 0
    while index < len(blocks):
        kind, text = blocks[index]
        safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind in {"heading", "subheading"}:
            style = heading if kind == "heading" else subheading
            display_text = safe_text.upper() if kind == "heading" else safe_text
            grouped: list[Flowable] = [Paragraph(display_text, style)]
            next_index = index + 1
            while next_index < len(blocks) and blocks[next_index][0] == "space":
                next_index += 1
            if next_index < len(blocks):
                next_kind, next_text = blocks[next_index]
                escaped_next = (
                    next_text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                if next_kind == "bullet":
                    grouped.append(Paragraph(escaped_next, bullet, bulletText="•"))
                    index = next_index
                elif next_kind == "body":
                    grouped.append(Paragraph(escaped_next, body))
                    index = next_index
            story.append(KeepTogether(grouped))
        elif kind == "space":
            story.append(Spacer(1, 2))
        elif kind == "title":
            story.append(Paragraph(safe_text, title))
        elif kind == "bullet":
            story.append(Paragraph(safe_text, bullet, bulletText="•"))
        else:
            story.append(Paragraph(safe_text, body))
        index += 1
    pdf = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Candidate-approved CV",
        author="GradPath AI",
    )
    pdf.build(story)
    return output.getvalue()
