"""Document boundary tests for import, export, and privacy controls."""

from io import BytesIO

import pytest
from docx import Document
from gradpath_api.core.config import Settings
from gradpath_api.main import create_app
from httpx import ASGITransport, AsyncClient
from pypdf import PdfReader


@pytest.mark.asyncio
async def test_extracts_docx_in_memory() -> None:
    document = Document()
    document.add_heading("Alex Example", level=1)
    document.add_paragraph(
        "Built and tested a TypeScript web application for graduates."
    )
    content = BytesIO()
    document.save(content)
    app = create_app(Settings(_env_file=None, environment="test"))

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.post(
            "/api/v1/documents/extract",
            files={
                "file": (
                    "cv.docx",
                    content.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    assert "TypeScript web application" in response.json()["text"]


@pytest.mark.asyncio
async def test_rejects_unsupported_document_type() -> None:
    app = create_app(Settings(_env_file=None, environment="test"))
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.post(
            "/api/v1/documents/extract",
            files={"file": ("cv.exe", b"This is not an accepted candidate document.")},
        )
    assert response.status_code == 422
    assert "TXT, Markdown, DOCX" in response.json()["detail"]


@pytest.mark.asyncio
async def test_exports_editable_docx_and_readable_pdf() -> None:
    app = create_app(Settings(_env_file=None, environment="test"))
    payload = {
        "filename": "alex-example-cv",
        "approved_cv_markdown": (
            "# Alex Example\n\n## Profile\n\nGraduate software developer.\n\n"
            "## Projects\n\n- Built a tested TypeScript application."
        ),
    }
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        docx_response = await client.post("/api/v1/documents/export/docx", json=payload)
        pdf_response = await client.post("/api/v1/documents/export/pdf", json=payload)

    assert docx_response.status_code == 200
    document = Document(BytesIO(docx_response.content))
    assert "Alex Example" in "\n".join(item.text for item in document.paragraphs)
    assert pdf_response.status_code == 200
    pdf = PdfReader(BytesIO(pdf_response.content))
    assert "Graduate software developer" in (pdf.pages[0].extract_text() or "")
